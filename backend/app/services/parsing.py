"""Extracts plain text (with light structural hints) from uploaded files.

Each parser returns a list of (text, heading) blocks — 'heading' is the
nearest section/heading title we could detect, used later to make citations
more meaningful than a bare page number. When no heading is detectable,
heading is None and the chunker falls back to plain windowing.
"""
import io
import re

from docx import Document as DocxDocument
from pypdf import PdfReader

HEADING_PATTERN = re.compile(
    r"^\s*("
    r"(chapter|section|clause|rule|regulation)\s+[\w.]+|"
    r"\d+(\.\d+){0,4}\s+[A-Z].{3,80}|"
    r"[A-Z][A-Z \-&/]{6,80}"
    r")\s*$"
)


def _is_heading(line: str) -> bool:
    line = line.strip()
    if not (3 < len(line) < 90):
        return False
    return bool(HEADING_PATTERN.match(line))


def parse_txt(raw: bytes) -> list[tuple[str, str | None]]:
    text = raw.decode("utf-8", errors="ignore")
    return _split_on_headings(text)


def parse_pdf(raw: bytes) -> list[tuple[str, str | None]]:
    reader = PdfReader(io.BytesIO(raw))
    blocks: list[tuple[str, str | None]] = []
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if not text.strip():
            continue
        for chunk_text, heading in _split_on_headings(text):
            label = heading or f"page {page_num}"
            blocks.append((chunk_text, label))
    return blocks


def parse_docx(raw: bytes) -> list[tuple[str, str | None]]:
    doc = DocxDocument(io.BytesIO(raw))
    blocks: list[tuple[str, str | None]] = []
    current_heading: str | None = None
    buffer: list[str] = []

    def flush():
        if buffer:
            blocks.append(("\n".join(buffer).strip(), current_heading))
            buffer.clear()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if "heading" in style or _is_heading(text):
            flush()
            current_heading = text
        else:
            buffer.append(text)
    flush()
    return blocks or [(p.text, None) for p in doc.paragraphs if p.text.strip()]


def _split_on_headings(text: str) -> list[tuple[str, str | None]]:
    lines = text.splitlines()
    blocks: list[tuple[str, str | None]] = []
    current_heading: str | None = None
    buffer: list[str] = []

    def flush():
        joined = "\n".join(buffer).strip()
        if joined:
            blocks.append((joined, current_heading))
        buffer.clear()

    for line in lines:
        if _is_heading(line):
            flush()
            current_heading = line.strip()
        else:
            buffer.append(line)
    flush()
    return blocks or [(text.strip(), None)]


PARSERS = {
    "pdf": parse_pdf,
    "docx": parse_docx,
    "txt": parse_txt,
}


def extract(file_type: str, raw: bytes) -> list[tuple[str, str | None]]:
    parser = PARSERS.get(file_type)
    if not parser:
        raise ValueError(f"Unsupported file type: {file_type}")
    return parser(raw)
